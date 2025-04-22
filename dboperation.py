import pyodbc
from docx import Document
from docx.shared import Pt

# ─── CONFIG ────────────────────────────────────────────────────────────────
SERVER = "DevSQL.msmyid.uga.edu"
USER = "python_analytics"
PASSWORD = "python_analytics"
DRIVER = "{ODBC Driver 17 for SQL Server}"
SYSTEM_DBS = ("master", "tempdb", "model", "msdb")

# Quick “recipes” for tables/views
QUERY_RECIPES = {
    "Top 10 Rows": "SELECT TOP 10 * FROM [{schema}].[{table}];",
    "Total Row Count": "SELECT COUNT(*) AS cnt FROM [{schema}].[{table}];",
}

# ─── HELPERS ───────────────────────────────────────────────────────────────


def get_connection(db_name=None):
    parts = [f"DRIVER={DRIVER}", f"SERVER={SERVER}", f"UID={USER}", f"PWD={PASSWORD}"]
    if db_name:
        parts.append(f"DATABASE={db_name}")
    return pyodbc.connect(";".join(parts))


def list_databases():
    """Connect to master and return all non-system databases."""
    with get_connection("master").cursor() as cur:
        cur.execute(
            """
            SELECT name
              FROM sys.databases
             WHERE state = 0
               AND name NOT IN (?,?,?,?)
             ORDER BY name
        """,
            SYSTEM_DBS,
        )
        return [row.name for row in cur.fetchall()]


def list_tables(cursor):
    cursor.execute(
        """
        SELECT TABLE_SCHEMA, TABLE_NAME
          FROM INFORMATION_SCHEMA.TABLES
         WHERE TABLE_TYPE = 'BASE TABLE'
         ORDER BY TABLE_SCHEMA, TABLE_NAME
    """
    )
    return cursor.fetchall()


def list_views(cursor):
    cursor.execute(
        """
        SELECT TABLE_SCHEMA, TABLE_NAME
          FROM INFORMATION_SCHEMA.VIEWS
         ORDER BY TABLE_SCHEMA, TABLE_NAME
    """
    )
    return cursor.fetchall()


def list_columns(cursor, schema, obj_name):
    cursor.execute(
        """
        SELECT COLUMN_NAME, DATA_TYPE, IS_NULLABLE, CHARACTER_MAXIMUM_LENGTH
          FROM INFORMATION_SCHEMA.COLUMNS
         WHERE TABLE_SCHEMA = ? 
           AND TABLE_NAME  = ?
         ORDER BY ORDINAL_POSITION
    """,
        (schema, obj_name),
    )
    return cursor.fetchall()


# ─── MAIN: build the doc ───────────────────────────────────────────────────


def make_full_db_cookbook(output_path="Full_Database_Cookbook.docx"):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # Title page
    doc.add_heading("Full SQL Server Cookbook", level=0)
    doc.add_paragraph(f"Server: {SERVER}")
    doc.add_page_break()

    # Enumerate all your user DBs
    for db in list_databases():
        doc.add_heading(db, level=1)
        conn = get_connection(db)
        cur = conn.cursor()

        tables = list_tables(cur)
        views = list_views(cur)

        # — Tables —
        if tables:
            doc.add_heading("Tables", level=2)
            tbl = doc.add_table(rows=1, cols=5)
            hdr = tbl.rows[0].cells
            hdr[0].text = "Schema"
            hdr[1].text = "Table"
            hdr[2].text = "Column"
            hdr[3].text = "Data Type"
            hdr[4].text = "Nullable"

            for schema, table in tables:
                for col_name, data_type, is_nullable, max_len in list_columns(
                    cur, schema, table
                ):
                    row = tbl.add_row().cells
                    row[0].text = schema
                    row[1].text = table
                    row[2].text = col_name
                    row[3].text = f"{data_type}" + (f"({max_len})" if max_len else "")
                    row[4].text = is_nullable
            doc.add_paragraph()

            # ─ Table recipes ─
            doc.add_heading("Table Recipes", level=3)
            for schema, table in tables:
                doc.add_paragraph(f"{schema}.{table}", style="List Bullet")
                for title, q in QUERY_RECIPES.items():
                    p = doc.add_paragraph(f"  • {title}: ")
                    # ← FIXED: use table=table, not object=table
                    p.add_run(q.format(schema=schema, table=table)).italic = True
                doc.add_paragraph()

        # — Views —
        if views:
            doc.add_heading("Views", level=2)
            for schema, view in views:
                doc.add_heading(f"{schema}.{view}", level=3)
                cols = list_columns(cur, schema, view)
                if not cols:
                    doc.add_paragraph("_No columns found for this view._")
                else:
                    vtbl = doc.add_table(rows=1, cols=4)
                    vh = vtbl.rows[0].cells
                    vh[0].text = "Column"
                    vh[1].text = "Data Type"
                    vh[2].text = "Nullable"
                    vh[3].text = "Max Length"
                    for col_name, data_type, is_nullable, max_len in cols:
                        r = vtbl.add_row().cells
                        r[0].text = col_name
                        r[1].text = data_type
                        r[2].text = is_nullable
                        r[3].text = str(max_len or "")
                doc.add_paragraph()

                # ─ View recipes ─
                doc.add_paragraph("Sample Query:", style="List Number")
                p = doc.add_paragraph()
                # ← FIXED: use table=view
                p.add_run(
                    QUERY_RECIPES["Top 10 Rows"].format(schema=schema, table=view)
                ).italic = True
                doc.add_paragraph()

        # If truly empty
        if not tables and not views:
            doc.add_paragraph("_No user tables or views found in this database._")

        conn.close()
        doc.add_page_break()

    doc.save(output_path)
    print(f"✅ Wrote cookbook (with tables & views) to {output_path}")


if __name__ == "__main__":
    make_full_db_cookbook()
